"""
Document Processor - Processamento de Documentos
Extrai texto de múltiplos formatos (PDF, DOCX, TXT)
"""

import os
import logging
from typing import Dict, List, Optional
from pathlib import Path
import mimetypes

# Configurar logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Processador de documentos multi-formato
    Suporta PDF, DOCX, TXT e outros formatos
    """
    
    def __init__(self):
        """Inicializa o processador de documentos"""
        self.supported_formats = {
            '.pdf': self._extract_pdf,
            '.docx': self._extract_docx,
            '.txt': self._extract_txt,
            '.md': self._extract_txt,
            '.rtf': self._extract_rtf
        }
        logger.info("📄 Document Processor inicializado")
    
    def extract_text(self, file_path: str) -> Dict:
        """
        Extrai texto de um arquivo
        
        Args:
            file_path: Caminho para o arquivo
            
        Returns:
            Dict: Dados extraídos do documento
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Arquivo não encontrado: {file_path}")
        
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Formato não suportado: {file_extension}")
        
        logger.info(f"📖 Extraindo texto de: {file_path}")
        
        try:
            # Chamar função de extração apropriada
            extract_function = self.supported_formats[file_extension]
            result = extract_function(file_path)
            
            # Adicionar metadados
            result.update({
                'file_path': file_path,
                'file_name': Path(file_path).name,
                'file_extension': file_extension,
                'file_size': os.path.getsize(file_path)
            })
            
            logger.info(f"✅ Texto extraído: {len(result['content'])} caracteres")
            return result
            
        except Exception as e:
            logger.error(f"❌ Erro extraindo texto de {file_path}: {e}")
            raise
    
    def _extract_pdf(self, file_path: str) -> Dict:
        """Extrai texto de arquivo PDF"""
        try:
            import PyPDF2
            
            content = ""
            metadata = {}
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extrair metadados
                if pdf_reader.metadata:
                    metadata = {
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'creator': pdf_reader.metadata.get('/Creator', ''),
                        'pages': len(pdf_reader.pages)
                    }
                
                # Extrair texto de todas as páginas
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        content += f"\n--- Página {page_num + 1} ---\n{page_text}\n"
                    except Exception as e:
                        logger.warning(f"⚠️ Erro na página {page_num + 1}: {e}")
            
            return {
                'content': content.strip(),
                'metadata': metadata,
                'format': 'pdf'
            }
            
        except ImportError:
            logger.error("❌ PyPDF2 não instalado. Use: pip install PyPDF2")
            return {'content': f"Erro: PyPDF2 não encontrado para processar {file_path}", 'metadata': {}, 'format': 'pdf'}
        except Exception as e:
            logger.error(f"❌ Erro processando PDF: {e}")
            return {'content': f"Erro processando PDF: {e}", 'metadata': {}, 'format': 'pdf'}
    
    def _extract_docx(self, file_path: str) -> Dict:
        """Extrai texto de arquivo DOCX"""
        try:
            import docx
            
            doc = docx.Document(file_path)
            content = ""
            
            # Extrair texto de parágrafos
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            # Extrair texto de tabelas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        content += cell.text + " "
                    content += "\n"
            
            # Metadados básicos
            metadata = {
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables)
            }
            
            return {
                'content': content.strip(),
                'metadata': metadata,
                'format': 'docx'
            }
            
        except ImportError:
            logger.error("❌ python-docx não instalado. Use: pip install python-docx")
            return {'content': f"Erro: python-docx não encontrado para processar {file_path}", 'metadata': {}, 'format': 'docx'}
        except Exception as e:
            logger.error(f"❌ Erro processando DOCX: {e}")
            return {'content': f"Erro processando DOCX: {e}", 'metadata': {}, 'format': 'docx'}
    
    def _extract_txt(self, file_path: str) -> Dict:
        """Extrai texto de arquivo TXT/MD"""
        try:
            # Tentar diferentes encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            content = ""
            encoding_used = ""
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                    encoding_used = encoding
                    break
                except UnicodeDecodeError:
                    continue
            
            if not content:
                raise ValueError("Não foi possível decodificar o arquivo com nenhum encoding")
            
            # Metadados básicos
            lines = content.split('\n')
            metadata = {
                'lines': len(lines),
                'encoding': encoding_used,
                'chars': len(content)
            }
            
            return {
                'content': content.strip(),
                'metadata': metadata,
                'format': 'txt'
            }
            
        except Exception as e:
            logger.error(f"❌ Erro processando TXT: {e}")
            return {'content': f"Erro processando arquivo de texto: {e}", 'metadata': {}, 'format': 'txt'}
    
    def _extract_rtf(self, file_path: str) -> Dict:
        """Extrai texto de arquivo RTF"""
        try:
            from striprtf.striprtf import rtf_to_text
            
            with open(file_path, 'r', encoding='utf-8') as file:
                rtf_content = file.read()
            
            content = rtf_to_text(rtf_content)
            
            metadata = {
                'chars': len(content)
            }
            
            return {
                'content': content.strip(),
                'metadata': metadata,
                'format': 'rtf'
            }
            
        except ImportError:
            logger.error("❌ striprtf não instalado. Use: pip install striprtf")
            return {'content': f"Erro: striprtf não encontrado para processar {file_path}", 'metadata': {}, 'format': 'rtf'}
        except Exception as e:
            logger.error(f"❌ Erro processando RTF: {e}")
            return {'content': f"Erro processando RTF: {e}", 'metadata': {}, 'format': 'rtf'}
    
    def process_directory(self, directory_path: str, recursive: bool = True) -> List[Dict]:
        """
        Processa todos os documentos em um diretório
        
        Args:
            directory_path: Caminho do diretório
            recursive: Se deve processar subdiretórios
            
        Returns:
            List[Dict]: Lista de documentos processados
        """
        if not os.path.exists(directory_path):
            raise FileNotFoundError(f"Diretório não encontrado: {directory_path}")
        
        documents = []
        path_obj = Path(directory_path)
        
        # Padrão de busca
        pattern = "**/*" if recursive else "*"
        
        for file_path in path_obj.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_formats:
                try:
                    doc_data = self.extract_text(str(file_path))
                    documents.append(doc_data)
                except Exception as e:
                    logger.error(f"❌ Erro processando {file_path}: {e}")
        
        logger.info(f"📁 Processados {len(documents)} documentos de {directory_path}")
        return documents
    
    def validate_file(self, file_path: str) -> Dict:
        """
        Valida se um arquivo pode ser processado
        
        Args:
            file_path: Caminho do arquivo
            
        Returns:
            Dict: Resultado da validação
        """
        result = {
            'is_valid': False,
            'errors': [],
            'warnings': [],
            'file_info': None
        }
        
        try:
            # Verificar se arquivo existe
            if not os.path.exists(file_path):
                result['errors'].append("Arquivo não encontrado")
                return result
            
            # Obter informações do arquivo
            file_info = self.get_document_info(file_path)
            result['file_info'] = file_info
            
            # Verificar formato suportado
            if not file_info['is_supported']:
                result['errors'].append(f"Formato não suportado: {file_info['file_extension']}")
                return result
            
            # Verificar tamanho do arquivo (limite de 100MB)
            if file_info['file_size'] > 100 * 1024 * 1024:
                result['errors'].append("Arquivo muito grande (máximo 100MB)")
                return result
            
            # Avisos para arquivos grandes
            if file_info['file_size'] > 10 * 1024 * 1024:
                result['warnings'].append("Arquivo grande, processamento pode demorar")
            
            result['is_valid'] = True
            
        except Exception as e:
            result['errors'].append(f"Erro na validação: {e}")
        
        return result
    
    def get_document_info(self, file_path: str) -> Dict:
        """
        Obtém informações básicas do documento sem extrair todo o texto
        
        Args:
            file_path: Caminho do arquivo
            
        Returns:
            Dict: Informações do documento
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Arquivo não encontrado: {file_path}")
        
        file_stats = os.stat(file_path)
        file_path_obj = Path(file_path)
        
        # Detectar tipo MIME
        mime_type, _ = mimetypes.guess_type(file_path)
        
        info = {
            'file_name': file_path_obj.name,
            'file_path': file_path,
            'file_extension': file_path_obj.suffix.lower(),
            'file_size': file_stats.st_size,
            'file_size_mb': round(file_stats.st_size / (1024 * 1024), 2),
            'mime_type': mime_type,
            'created_time': file_stats.st_ctime,
            'modified_time': file_stats.st_mtime,
            'is_supported': file_path_obj.suffix.lower() in self.supported_formats
        }
        
        return info

def main():
    """Função de teste"""
    processor = DocumentProcessor()
    
    # Testar com arquivo de exemplo
    test_files = [
        "data/raw/sample.txt",
        "data/raw/sample.pdf", 
        "data/raw/sample.docx"
    ]
    
    for file_path in test_files:
        if os.path.exists(file_path):
            try:
                result = processor.extract_text(file_path)
                print(f"✅ {file_path}: {len(result['content'])} caracteres")
            except Exception as e:
                print(f"❌ {file_path}: {e}")
        else:
            print(f"⚠️ Arquivo não encontrado: {file_path}")

if __name__ == "__main__":
    main()